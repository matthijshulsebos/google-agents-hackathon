"""
File ingestion pipeline for extracting text from multiple file formats.
"""
import io
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import pdfplumber
from docx import Document
import openpyxl
import pandas as pd
from bs4 import BeautifulSoup
from google.cloud import storage
import tiktoken

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various file formats."""
    
    @staticmethod
    def extract_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file."""
        text = []
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
        return "\n\n".join(text)
    
    @staticmethod
    def extract_from_docx(file_content: bytes) -> str:
        """Extract text from Word document."""
        text = []
        try:
            doc = Document(io.BytesIO(file_content))
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text.append(row_text)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
        return "\n\n".join(text)
    
    @staticmethod
    def extract_from_xlsx(file_content: bytes) -> str:
        """Extract text from Excel file."""
        text = []
        try:
            # Read all sheets
            df_dict = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
            for sheet_name, df in df_dict.items():
                text.append(f"=== Sheet: {sheet_name} ===")
                # Convert dataframe to string representation
                text.append(df.to_string(index=False))
        except Exception as e:
            logger.error(f"Error extracting XLSX: {e}")
        return "\n\n".join(text)
    
    @staticmethod
    def extract_from_csv(file_content: bytes) -> str:
        """Extract text from CSV file."""
        try:
            df = pd.read_csv(io.BytesIO(file_content))
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Error extracting CSV: {e}")
            return ""
    
    @staticmethod
    def extract_from_html(file_content: bytes) -> str:
        """Extract text from HTML file."""
        try:
            soup = BeautifulSoup(file_content, 'lxml')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            return text
        except Exception as e:
            logger.error(f"Error extracting HTML: {e}")
            return ""
    
    @staticmethod
    def extract_from_txt(file_content: bytes) -> str:
        """Extract text from plain text file."""
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_content.decode('latin-1')
            except Exception as e:
                logger.error(f"Error extracting TXT: {e}")
                return ""
    
    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text based on file type."""
        file_type = file_type.lower()
        
        extractors = {
            'pdf': self.extract_from_pdf,
            'docx': self.extract_from_docx,
            'doc': self.extract_from_docx,
            'xlsx': self.extract_from_xlsx,
            'xls': self.extract_from_xlsx,
            'csv': self.extract_from_csv,
            'html': self.extract_from_html,
            'htm': self.extract_from_html,
            'txt': self.extract_from_txt,
            'text': self.extract_from_txt,
        }
        
        extractor = extractors.get(file_type)
        if extractor:
            return extractor(file_content)
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return ""


class TextChunker:
    """Chunk text into smaller pieces for indexing."""
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 200, min_chunk_size: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except:
            self.tokenizer = None
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        if self.tokenizer:
            return len(self.tokenizer.encode(text))
        else:
            # Rough approximation: 1 token ~= 4 characters
            return len(text) // 4
    
    def chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces with overlap."""
        if not text or len(text.strip()) < self.min_chunk_size:
            return []
        
        chunks = []
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_tokens = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            para_tokens = self.count_tokens(paragraph)
            
            # If single paragraph exceeds chunk size, split by sentences
            if para_tokens > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                    current_tokens = 0
                
                # Split long paragraph by sentences
                sentences = paragraph.replace('! ', '!|').replace('? ', '?|').replace('. ', '.|').split('|')
                for sentence in sentences:
                    sent_tokens = self.count_tokens(sentence)
                    if current_tokens + sent_tokens > self.chunk_size and current_chunk:
                        chunks.append(current_chunk.strip())
                        # Keep overlap
                        overlap_text = current_chunk[-self.chunk_overlap*4:] if len(current_chunk) > self.chunk_overlap*4 else current_chunk
                        current_chunk = overlap_text + " " + sentence
                        current_tokens = self.count_tokens(current_chunk)
                    else:
                        current_chunk += " " + sentence
                        current_tokens += sent_tokens
            
            # Normal case: add paragraph to current chunk
            elif current_tokens + para_tokens > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                # Keep overlap
                overlap_text = current_chunk[-self.chunk_overlap*4:] if len(current_chunk) > self.chunk_overlap*4 else current_chunk
                current_chunk = overlap_text + "\n\n" + paragraph
                current_tokens = self.count_tokens(current_chunk)
            else:
                current_chunk += "\n\n" + paragraph
                current_tokens += para_tokens
        
        # Add remaining chunk
        if current_chunk and len(current_chunk.strip()) >= self.min_chunk_size:
            chunks.append(current_chunk.strip())
        
        return chunks


class DocumentProcessor:
    """Process documents from GCS buckets."""
    
    def __init__(self, project_id: str, chunk_size: int = 800, chunk_overlap: int = 200):
        self.project_id = project_id
        self.storage_client = storage.Client(project=project_id)
        self.text_extractor = TextExtractor()
        self.text_chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    
    def get_file_type(self, filename: str) -> str:
        """Get file extension from filename."""
        return Path(filename).suffix.lstrip('.').lower()
    
    def process_file(self, bucket_name: str, blob_name: str, domain: str) -> List[Dict[str, Any]]:
        """Process a single file from GCS."""
        try:
            bucket = self.storage_client.bucket(bucket_name)
            blob = bucket.blob(blob_name)
            
            # Download file content
            file_content = blob.download_as_bytes()
            file_type = self.get_file_type(blob_name)
            
            # Extract text
            text = self.text_extractor.extract_text(file_content, file_type)
            
            if not text:
                logger.warning(f"No text extracted from {blob_name}")
                return []
            
            # Chunk text
            chunks = self.text_chunker.chunk_text(text)
            
            # Create metadata for each chunk
            documents = []
            for idx, chunk in enumerate(chunks):
                doc = {
                    "content": chunk,
                    "metadata": {
                        "domain": domain,
                        "bucket": bucket_name,
                        "filename": blob_name,
                        "file_type": file_type,
                        "chunk_id": idx,
                        "total_chunks": len(chunks)
                    }
                }
                documents.append(doc)
            
            logger.info(f"Processed {blob_name}: {len(chunks)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing file {blob_name}: {e}")
            return []
    
    def process_bucket(self, bucket_name: str, domain: str) -> List[Dict[str, Any]]:
        """Process all files in a bucket."""
        all_documents = []
        
        try:
            bucket = self.storage_client.bucket(bucket_name)
            blobs = bucket.list_blobs()
            
            for blob in blobs:
                # Skip directories
                if blob.name.endswith('/'):
                    continue
                
                documents = self.process_file(bucket_name, blob.name, domain)
                all_documents.extend(documents)
            
            logger.info(f"Processed bucket {bucket_name}: {len(all_documents)} total chunks")
            
        except Exception as e:
            logger.error(f"Error processing bucket {bucket_name}: {e}")
        
        return all_documents
